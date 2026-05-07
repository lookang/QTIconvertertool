from docx import Document
from docx.oxml.ns import qn
import os
import re


# ─────────────────────────────────────────────────────────────────────────────
# MARK-SCHEME PARSER
# ─────────────────────────────────────────────────────────────────────────────

def parse_mark_scheme(ms_docx_path):
    """
    Parse a mark-scheme .docx and return a dict  {qnum_str: answer}.

    MCQ papers (P1):
        The first table is a 6-row × 10-col grid of alternating
        [question-numbers, answer-letters] rows. Each answer letter is stored
        as lowercase ('a'–'d') to match QTI simpleChoice identifiers.

    Structured papers (P2 / P3):
        The table rows have the form  [part_label, answer_text, marks…].
        All parts belonging to the same question number are concatenated.
        Result: { '1': '1a: …\n1bi: …\n…', '2': '2a: …', … }
    """
    doc    = Document(ms_docx_path)
    tables = doc.tables
    if not tables:
        return {}

    first  = tables[0]
    nrows  = len(first.rows)
    ncols  = len(first.columns) if first.rows else 0

    # --- MCQ grid detection ---
    if nrows >= 6 and ncols >= 10:
        first_cell = first.rows[0].cells[0].text.strip()
        if first_cell.isdigit():
            return _parse_mcq_grid(first)

    # --- Structured MS ---
    return _parse_structured_ms(tables)


def _normalize_mcq_answer(text):
    """
    Normalize a mark-scheme MCQ answer token to the QTI choice id.

    Examples:
        'A' -> 'a'
        '1' -> 'a'
        '(1)' -> 'a'
        '1)' -> 'a'
        '①' -> 'a'
        '⑴' -> 'a'
    """
    compact = re.sub(r'\s+', '', (text or '')).upper()
    if not compact:
        return None

    if compact in ("A", "B", "C", "D"):
        return compact.lower()

    circled_map = {
        "①": "1",
        "②": "2",
        "③": "3",
        "④": "4",
        "⑴": "1",
        "⑵": "2",
        "⑶": "3",
        "⑷": "4",
    }
    compact = circled_map.get(compact, compact)

    if compact.startswith(("(", "（")):
        compact = compact[1:]
    if compact.endswith((")", "）", ".", "．")):
        compact = compact[:-1]

    if compact in ("1", "2", "3", "4"):
        return chr(96 + int(compact))

    return None


def _parse_mcq_grid(grid_table):
    """
    Alternating rows: [q-numbers …], [answer-letters …], [q-numbers …], …
    Returns { '1': 'b', '2': 'c', … } (lowercase letters).
    """
    answers = {}
    rows = [[c.text.strip() for c in r.cells] for r in grid_table.rows]

    for i in range(0, len(rows) - 1, 2):
        nums_row = rows[i]
        ans_row  = rows[i + 1]
        for qn, ans in zip(nums_row, ans_row):
            qn  = qn.strip()
            ans = ans.strip().upper()
            if qn.isdigit() and ans in ('A', 'B', 'C', 'D'):
                answers[qn] = ans.lower()

    return answers


def _parse_structured_ms(all_tables):
    """
    Rows: [part_label, answer_text, marks…]
    e.g. ['1a', 'rate of change of velocity', 'B1']
         ['1bi', 'Taking upwards …', 'C1\nA1']
    Returns { '1': '1a: rate of change…\n1bi: Taking upwards…', … }
    """
    answers = {}

    for table in all_tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]

            paired_answers = 0
            for idx in range(0, len(cells) - 1, 2):
                qn = cells[idx].strip()
                ans = _normalize_mcq_answer(cells[idx + 1])
                if qn.isdigit() and ans:
                    answers[qn] = ans
                    paired_answers += 1

            if paired_answers:
                continue

            if len(cells) < 2:
                continue
            label = cells[0]
            text  = cells[1]
            if not label or not text:
                continue
            m = re.match(r'^(\d+)', label)
            if not m:
                continue
            qnum = m.group(1)
            entry = f"{label}: {text}"
            if qnum in answers:
                answers[qnum] += "\n" + entry
            else:
                answers[qnum] = entry

    return answers


# ─────────────────────────────────────────────────────────────────────────────
# MAIN CONVERTER
# ─────────────────────────────────────────────────────────────────────────────

def convert_docx_to_qti(input_docx, job_dir, ms_answers=None):
    """
    Convert a question-paper .docx to a QTI 2.1 ZIP package.

    Parameters
    ----------
    input_docx  : path to the question-paper .docx
    job_dir     : working directory for output files
    ms_answers  : optional dict from parse_mark_scheme(); keys are question
                  number strings ('1', '2', …); values are answer strings.
                  For MCQ items the value should be a single lowercase letter
                  ('a'–'d') matching the simpleChoice identifier.
    """
    doc = Document(input_docx)

    assets_dir = os.path.join(job_dir, "assets")
    items_dir  = os.path.join(job_dir, "items")
    os.makedirs(assets_dir, exist_ok=True)

    if ms_answers is None:
        ms_answers = {}

    # -----------------------------
    # IMAGE HANDLING
    # -----------------------------
    image_counter = 0

    def save_image(doc, rId):
        nonlocal image_counter
        image_part = doc.part.related_parts[rId]
        filename   = f"img_{image_counter}.jpg"
        with open(os.path.join(assets_dir, filename), "wb") as f:
            f.write(image_part.blob)
        image_counter += 1
        return filename

    # -----------------------------
    # ITERATE WORD CONTENT
    # -----------------------------

    def extract_run_text(run):
        """Extract text from a single run, applying sup/sub/bold/italic HTML tags."""
        text = "".join(node.text or "" for node in run.xpath(".//w:t"))
        if not text:
            return ""
        is_sup    = run.xpath(".//w:vertAlign[@w:val='superscript']")
        is_sub    = run.xpath(".//w:vertAlign[@w:val='subscript']")
        is_bold   = run.xpath(".//w:b")
        is_italic = run.xpath(".//w:i")
        if is_sup:    text = f"<sup>{text}</sup>"
        if is_sub:    text = f"<sub>{text}</sub>"
        if is_bold:   text = f"<strong>{text}</strong>"
        if is_italic: text = f"<em>{text}</em>"
        return text

    def extract_cell_text(cell):
        """Extract formatted text (sup/sub/bold/italic) from a table cell element."""
        return "".join(extract_run_text(r) for r in cell.xpath(".//w:r")).strip()

    def iter_block_items(doc):
        body = doc.element.body
        for child in body.iterchildren():

            # ── Paragraph ──────────────────────────────────────────────
            if child.tag.endswith('p'):

                paragraph_text = "".join(
                    extract_run_text(r) for r in child.xpath("./w:r")
                ).strip()

                if paragraph_text:
                    yield ("text", paragraph_text)

                for node in child.iter():
                    if node.tag.endswith('blip'):
                        rId = node.get(qn('r:embed'))
                        if rId:
                            yield ("image", rId)

            # ── Table ──────────────────────────────────────────────────
            if child.tag.endswith('tbl'):

                table_data      = []
                seen_image_rids = []   # deduplicate images within one table

                for row in child.xpath(".//w:tr"):

                    row_data = []

                    for cell in row.xpath(".//w:tc"):
                        # FIX: use formatted text (preserves <sup>/<sub>)
                        row_data.append(extract_cell_text(cell))

                        # FIX: collect images from table cells
                        for node in cell.iter():
                            if node.tag.endswith('blip'):
                                rId = node.get(qn('r:embed'))
                                if rId and rId not in seen_image_rids:
                                    seen_image_rids.append(rId)

                    table_data.append(row_data)

                yield ("table", table_data)

                # Yield images discovered inside table cells so they are
                # attached to the current question token list
                for rId in seen_image_rids:
                    yield ("image", rId)

    # -----------------------------
    # CELL DEDUPLICATION
    # Merged cells in Word tables repeat identical text — keep only first
    # occurrence in each row so question text isn't duplicated.
    # -----------------------------
    def dedup_cells(cells):
        seen  = set()
        parts = []
        for c in cells:
            c = c.strip()
            if c and c not in seen:
                seen.add(c)
                parts.append(c)
        return ' '.join(parts)

    # Options row: ['', 'A   text', 'B   text', 'C   text', 'D   text']
    OPT_CELL_RE  = re.compile(r'^([A-D])\s+(.*)', re.DOTALL | re.I)
    # Single-letter MCQ label that IS an option identifier (not part of a word)
    OPT_LABEL_RE = re.compile(r'^[A-D]$', re.I)
    NUMBERED_OPT_LABEL_RE = re.compile(
        r'^(?:[\u2474-\u247B]|[\u2460-\u2467]|[\(\uff08]\s*[1-8]\s*[\)\uff09]|[1-8][\.\)\uff0e])$',
        re.I
    )

    def plain_cell_text(cell):
        return re.sub(r'<[^>]+>', '', cell or '').replace('\xa0', ' ').strip()

    def extract_opts_from_row(row):
        """Return option texts from a table row."""
        opts = []
        for cell in row[1:]:          # skip first (empty) cell
            cell = cell.strip()
            if not cell:
                continue
            m = OPT_CELL_RE.match(cell)
            if m:
                opts.append(m.group(2).strip())
            # lone 'A'/'B' etc (image-only options) → skip (image captured separately)
        if opts:
            return opts

        if row and not plain_cell_text(row[0]):
            label_idx = None
            for idx, cell in enumerate(row[1:], start=1):
                plain = plain_cell_text(cell)
                if not plain:
                    continue
                if not NUMBERED_OPT_LABEL_RE.match(plain):
                    label_idx = None
                    break
                label_idx = idx
                break

            if label_idx is not None:
                opt_text = dedup_cells(row[label_idx + 1:]).strip()
                if opt_text and not is_answer_blank(opt_text):
                    return [opt_text]

        return opts

    # -----------------------------
    # REGEX PATTERNS (STRICT)
    # -----------------------------
    question_regex  = re.compile(r'^(\d+)\s+(.*)')
    option_regex    = re.compile(
        r'^(?:[A-D][\.\)]|[\(\uff08]\s*[1-8]\s*[\)\uff09]|[1-8][\.\)\uff0e])\s*(.*)'
    )
    # Table first-cell is a bare question number, optionally with a short
    # non-digit prefix from Word cross-reference field codes (e.g. "XX23").
    qnum_only_regex = re.compile(r'^[A-Z]{0,5}(\d{1,3})$', re.I)
    answer_blank_regex = re.compile(r'^[\(\（][\s_＿]*[\)\）]$')
    inline_option_regex = re.compile(
        r'Q(?P<qnum>\d+)[_＿\s]*[（(]\s*1\s*(?P<a>.*?)\s*2\s*(?P<b>.*?)\s*3\s*(?P<c>.*?)\s*4\s*(?P<d>.*?)\s*[）)]',
        re.DOTALL
    )
    inline_option_strip_regex = re.compile(
        r'(Q\d+[_＿\s]*)[（(]\s*1\s*.*?\s*2\s*.*?\s*3\s*.*?\s*4\s*.*?\s*[）)]',
        re.DOTALL
    )

    MAX_QNUM = 200   # ignore "question numbers" > this (e.g. the year 2022)

    def parse_qnum(first_cell):
        """Return int qnum from first_cell if it looks like a question-start,
        else return None.  Handles bare '25' and prefixed 'XX23' alike."""
        normalized = re.sub(r'\s+', '', first_cell or '')
        m = qnum_only_regex.match(normalized)
        if m:
            n = int(m.group(1))
            if n <= MAX_QNUM:
                return str(n)   # normalise to plain digit string
        return None

    def is_answer_blank(text):
        compact = text.replace('\u3000', ' ').strip()
        return bool(compact) and bool(answer_blank_regex.match(compact))

    def find_qnum_in_row(row):
        for idx, cell in enumerate(row):
            plain = plain_cell_text(cell)
            if not plain:
                continue
            if OPT_LABEL_RE.match(plain) or NUMBERED_OPT_LABEL_RE.match(plain):
                return None, None
            qnum = parse_qnum(plain)
            if qnum is not None:
                return qnum, idx
            return None, None
        return None, None

    def extract_unlabeled_option(row):
        if not row or row[0].strip():
            return None
        unique_tail = []
        seen = set()
        for cell in row[1:]:
            cell = cell.strip()
            if not cell or cell in seen:
                continue
            seen.add(cell)
            unique_tail.append(cell)
        if len(unique_tail) != 1:
            return None
        candidate = unique_tail[0]
        if is_answer_blank(candidate) or parse_qnum(candidate) is not None:
            return None
        return candidate

    def extract_inline_options_from_tokens(tokens, qnum):
        joined = "\n".join(val for ttype, val in tokens if ttype == "text")
        if not joined:
            return []
        for match in inline_option_regex.finditer(joined):
            if match.group("qnum") != str(qnum):
                continue
            extracted = []
            for key in ("a", "b", "c", "d"):
                opt = re.sub(r'\s+', ' ', match.group(key)).strip()
                if opt:
                    extracted.append(opt)
            if len(extracted) >= 2:
                return extracted
        return []

    def strip_inline_options_from_tokens(tokens):
        cleaned = []
        for ttype, val in tokens:
            if ttype != "text":
                cleaned.append((ttype, val))
                continue
            stripped = inline_option_strip_regex.sub(r'\1', val)
            stripped = re.sub(r'\s{2,}', ' ', stripped).strip()
            if stripped:
                cleaned.append((ttype, stripped))
        return cleaned

    # -----------------------------
    # STORAGE
    # -----------------------------
    mcq_questions        = []
    structured_questions = []

    current_qnum                 = None
    current_tokens               = []
    options                      = []
    allow_implicit_table_options = False
    pending_lead_tokens          = []

    answers              = {}
    reading_answers      = False
    current_answer_q     = None
    current_answer_tokens = []

    def start_question(qnum, seed_tokens=None, allow_implicit_options=False):
        nonlocal current_qnum, current_tokens, options
        nonlocal allow_implicit_table_options, pending_lead_tokens
        current_qnum = qnum
        current_tokens = list(pending_lead_tokens)
        current_tokens.extend(seed_tokens or [])
        options = []
        allow_implicit_table_options = allow_implicit_options
        pending_lead_tokens = []

    def flush_current():
        if current_qnum is None:
            return
        inline_options = []
        tokens_for_output = current_tokens
        if not options:
            inline_options = extract_inline_options_from_tokens(current_tokens, current_qnum)
            if inline_options:
                tokens_for_output = strip_inline_options_from_tokens(current_tokens)
        if options or inline_options:
            mcq_questions.append({
                "qnum":    current_qnum,
                "tokens":  tokens_for_output,
                "options": options or inline_options
            })
        else:
            structured_questions.append({
                "qnum":   current_qnum,
                "tokens": current_tokens
            })

    # -----------------------------
    # PARSE DOCUMENT
    # -----------------------------
    for item_type, value in iter_block_items(doc):

        if reading_answers:

            if item_type == "image" and current_answer_q is not None:
                filename = save_image(doc, value)
                current_answer_tokens.append(("image", filename))
                continue

            if item_type == "table":
                current_answer_tokens.append(("table", value))
                continue

            if item_type == "text":

                text   = value.replace('\xa0', ' ').strip()
                amatch = re.match(r'^(\d+)', text)

                if amatch:

                    if current_answer_q is not None:
                        answers[current_answer_q] = current_answer_tokens

                    current_answer_q      = amatch.group(1)
                    cleaned               = re.sub(r'^\d+\s*', '', text)
                    current_answer_tokens = []

                    if cleaned:
                        current_answer_tokens.append(("text", cleaned))

                else:

                    if current_answer_q is not None:
                        current_answer_tokens.append(("text", text))

            continue

        if item_type == "image":
            filename = save_image(doc, value)
            current_tokens.append(("image", filename))
            continue

        if item_type == "table":
            rows = value
            if not rows:
                continue

            shared_table_tokens = []

            for row in rows:
                if not row:
                    continue

                detected_qnum, qnum_idx = find_qnum_in_row(row)
                if detected_qnum is not None:
                    flush_current()
                    q_text = dedup_cells(row[qnum_idx + 1:])
                    allow_implicit_options = any(
                        is_answer_blank(cell) for cell in row[qnum_idx + 1:]
                    )
                    start_question(
                        detected_qnum,
                        seed_tokens=shared_table_tokens,
                        allow_implicit_options=allow_implicit_options
                    )
                    if q_text and not is_answer_blank(q_text):
                        current_tokens.append(("text", q_text))
                    continue

                row0 = row[0].strip()

                if current_qnum is not None and row0 and OPT_LABEL_RE.match(row0):
                    opt_text = dedup_cells(row[1:])
                    if opt_text:
                        options.append(opt_text)
                    continue

                if current_qnum is not None and allow_implicit_table_options:
                    implicit_opt = extract_unlabeled_option(row)
                    if implicit_opt and len(options) < 4:
                        options.append(implicit_opt)
                        continue

                if current_qnum is not None and not row0:
                    opts = extract_opts_from_row(row)
                    if opts:
                        options.extend(opts)
                        continue

                extra = dedup_cells(row)
                if not extra:
                    continue

                current_has_inline_options = (
                    current_qnum is not None and
                    bool(extract_inline_options_from_tokens(current_tokens, current_qnum))
                )

                # Standalone rows that appear after an MCQ already has choices
                # (either explicit table options or inline numbered choices)
                # are typically instructions or shared passage text for the
                # next question block, not part of the previous item.
                if current_qnum is None or options or current_has_inline_options:
                    shared_table_tokens.append(("text", extra))
                else:
                    current_tokens.append(("text", extra))

            continue

        if item_type == "text":

            text = value.replace('\xa0', ' ').replace('\t', ' ')
            text = re.sub(r'\s+', ' ', text).strip()

            if text.upper() == "ANSWERS":
                reading_answers = True
                continue

            if not text:
                continue

            qmatch = question_regex.match(text)

            if qmatch and int(qmatch.group(1)) <= MAX_QNUM:
                flush_current()
                start_question(qmatch.group(1))
                current_tokens.append(("text", qmatch.group(2)))
                continue

            opt = option_regex.match(text)

            if opt:
                options.append(opt.group(1))
                continue

            current_has_inline_options = (
                current_qnum is not None and
                bool(extract_inline_options_from_tokens(current_tokens, current_qnum))
            )

            if current_qnum is not None and (options or current_has_inline_options):
                pending_lead_tokens.append(("text", text))
                continue

            current_tokens.append(("text", text))

    if current_answer_q is not None:
        answers[current_answer_q] = current_answer_tokens

    flush_current()

    # ─────────────────────────────────────────────────────────────────
    # HELPER: resolve answer for a question number
    # Priority: ms_answers (external MS file) > inline answers section
    # ─────────────────────────────────────────────────────────────────
    def resolve_answer(qnum, is_mcq=False):
        """Return answer string for the given question number."""
        # External mark-scheme takes priority
        if qnum in ms_answers:
            answer = ms_answers[qnum]
            if is_mcq:
                return _normalize_mcq_answer(answer) or answer
            return answer
        # Inline ANSWERS section
        if qnum in answers:
            answer = answer_tokens_to_string(answers[qnum])
            if is_mcq:
                return _normalize_mcq_answer(answer) or answer
            return answer
        return "None"

    # -----------------------------
    # QTI 2.1 OUTPUT GENERATION
    # -----------------------------
    import zipfile
    import random
    import string
    import xml.sax.saxutils as saxutils

    os.makedirs(items_dir, exist_ok=True)

    def random_id(length=3):
        return ''.join(random.choices(string.ascii_uppercase + string.digits, k=length))

    item_refs   = []
    item_images = {}

    # -----------------------------
    # HELPER: convert tokens to html
    # -----------------------------
    def tokens_to_html(tokens):

        html = ""

        for ttype, val in tokens:

            if ttype == "text":

                safe  = saxutils.escape(val)
                safe  = safe.replace("&lt;", "<").replace("&gt;", ">")
                html += f"<div>{safe}</div>\n"

            elif ttype == "image":

                html += f'''
<div>
<img alt="diagram" src="../assets/{val}"/>
</div>
'''

            elif ttype == "table":

                html += "<table border='1'>\n"
                for row in val:
                    html += "<tr>"
                    for cell in row:
                        safe  = saxutils.escape(cell)
                        html += f"<td>{safe}</td>"
                    html += "</tr>\n"
                html += "</table>\n"

        return html

    # -----------------------------
    # CONVERT ANSWER TOKENS TO STRING
    # -----------------------------
    def answer_tokens_to_string(tokens):

        result = ""

        for ttype, val in tokens:
            if ttype == "text":
                result += val + " "
            elif ttype == "image":
                result += f"[image:{val}] "
            elif ttype == "table":
                result += "[table] "

        return result.strip()

    letters = ["a", "b", "c", "d", "e", "f"]

    # -----------------------------
    # MCQ QUESTIONS → QTI
    # -----------------------------
    for q in mcq_questions:

        answer_text = resolve_answer(q["qnum"], is_mcq=True)

        item_id  = f"Q{int(q['qnum']):03d}_{random_id()}"
        filename = os.path.join(items_dir, f"{item_id}.xml")

        stem_html = tokens_to_html(q["tokens"])

        title_text = ""
        title_candidates = []
        for ttype, val in q["tokens"]:
            if ttype == "text":
                plain = re.sub(r'<[^>]+>', '', val).strip()
                if plain:
                    title_candidates.append(plain)

        for candidate in title_candidates:
            if any(marker in candidate for marker in ("______", "＿＿＿", "(      )", "(       )")):
                title_text = candidate[:70]
                break

        if not title_text:
            for candidate in title_candidates:
                if any(marker in candidate for marker in ("?", "？")):
                    title_text = candidate[:70]
                    break

        if not title_text and title_candidates:
            title_text = title_candidates[-1][:70]

        safe_title = saxutils.escape(title_text)

        xml = f'''<assessmentItem xmlns="http://www.imsglobal.org/xsd/imsqti_v2p1"
xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
adaptive="false"
identifier="{item_id}"
timeDependent="false"
title="Q{int(q['qnum']):03d} {safe_title}"
xsi:schemaLocation="http://www.imsglobal.org/xsd/imsqti_v2p1 http://www.imsglobal.org/xsd/qti/qtiv2p1/imsqti_v2p1.xsd">
<responseDeclaration identifier="RESPONSE" cardinality="single" baseType="identifier">
<correctResponse>
<value>{saxutils.escape(answer_text)}</value>
</correctResponse>
</responseDeclaration>
<outcomeDeclaration baseType="float" cardinality="single" identifier="SCORE"/>
<itemBody>
<choiceInteraction responseIdentifier="RESPONSE" shuffle="true" maxChoices="1">
<prompt>
<div>
{stem_html}
</div>
</prompt>
'''

        for i, opt in enumerate(q["options"]):
            if i >= len(letters):
                break
            safe = saxutils.escape(opt)
            xml += f'<simpleChoice identifier="{letters[i]}">{safe}</simpleChoice>\n'

        xml += '''
</choiceInteraction>
</itemBody>
<responseProcessing template="http://www.imsglobal.org/question/qti_v2p1/rptemplates/match_correct"/>
</assessmentItem>
'''

        with open(filename, "w", encoding="utf-8") as f:
            f.write(xml)

        item_refs.append(item_id)
        imgs = [v for t, v in q["tokens"] if t == "image"]
        item_images[item_id] = imgs

    # -----------------------------
    # STRUCTURED QUESTIONS → QTI
    # -----------------------------
    for q in structured_questions:

        answer_text = resolve_answer(q["qnum"])

        item_id  = f"Q{int(q['qnum']):03d}_{random_id()}"
        filename = os.path.join(items_dir, f"{item_id}.xml")

        stem_html = tokens_to_html(q["tokens"])

        xml = f'''<assessmentItem xmlns="http://www.imsglobal.org/xsd/imsqti_v2p1"
xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
adaptive="false"
identifier="{item_id}"
timeDependent="false"
title="Q{int(q['qnum']):03d}"
xsi:schemaLocation="http://www.imsglobal.org/xsd/imsqti_v2p1 http://www.imsglobal.org/xsd/qti/qtiv2p1/imsqti_v2p1.xsd">
<responseDeclaration identifier="RESPONSE" cardinality="single" baseType="string"/>
<correctResponse>
<value>{saxutils.escape(answer_text)}</value>
</correctResponse>
<outcomeDeclaration baseType="float" cardinality="single" identifier="SCORE"/>
<itemBody>
<prompt>
{stem_html}
</prompt>
<extendedTextInteraction responseIdentifier="RESPONSE" expectedLength="400"/>
</itemBody>
<responseProcessing template="http://www.imsglobal.org/question/qti_v2p1/rptemplates/match_correct"/>
</assessmentItem>
'''

        with open(filename, "w", encoding="utf-8") as f:
            f.write(xml)

        item_refs.append(item_id)
        imgs = [v for t, v in q["tokens"] if t == "image"]
        item_images[item_id] = imgs

    # -----------------------------
    # assessment_test.xml
    # -----------------------------
    assessment_xml = '''
<assessmentTest xmlns="http://www.imsglobal.org/xsd/imsqti_v2p1"
xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
identifier="TEST1" title="Converted Test" xsi:schemaLocation="http://www.imsglobal.org/xsd/imsqti_v2p1 http://www.imsglobal.org/xsd/qti/qtiv2p1/imsqti_v2p1.xsd">
<testPart identifier="part1" navigationMode="linear" submissionMode="individual">
<assessmentSection identifier="section1" title="Converted Test" visible="true">
'''

    for item in item_refs:
        assessment_xml += f'<assessmentItemRef href="items/{item}.xml" identifier="{item}"/>\n'

    assessment_xml += '''
</assessmentSection>
</testPart>
</assessmentTest>
'''

    with open(os.path.join(job_dir, "assessment_test.xml"), "w") as f:
        f.write(assessment_xml)

    # -----------------------------
    # imsmanifest.xml
    # -----------------------------
    manifest = '''
<manifest
xmlns="http://www.imsglobal.org/xsd/imscp_v1p1"
xmlns:imsmd="http://www.imsglobal.org/xsd/imsmd_v1p2"
xmlns:imsqti="http://www.imsglobal.org/xsd/imsqti_metadata_v2p1"
xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
identifier="MANIFEST1"
xsi:schemaLocation="
http://www.imsglobal.org/xsd/imscp_v1p1 http://www.imsglobal.org/xsd/imscp_v1p1.xsd
http://www.imsglobal.org/xsd/imsmd_v1p2 http://www.imsglobal.org/xsd/imsmd_v1p2p4.xsd
http://www.imsglobal.org/xsd/imsqti_metadata_v2p1 http://www.imsglobal.org/xsd/qti/qtiv2p1/imsqti_metadata_v2p1.xsd">
<metadata>
<schema>QTIv2.1 Package</schema>
<schemaversion>1.0.0</schemaversion>
</metadata>
<organizations/>
<resources>
'''

    manifest += '''
<resource identifier="RES_TEST" type="imsqti_test_xmlv2p1" href="assessment_test.xml">
<metadata/>
<file href="assessment_test.xml"/>
</resource>
'''

    for item in item_refs:

        manifest += f'''
<resource identifier="RES_{item}" type="imsqti_item_xmlv2p1" href="items/{item}.xml">
<metadata/>
<file href="items/{item}.xml"/>
'''

        for img in item_images.get(item, []):
            manifest += f'    <file href="assets/{img}"/>\n'

        manifest += "</resource>\n"

    manifest += '''
</resources>
</manifest>
'''

    with open(os.path.join(job_dir, "imsmanifest.xml"), "w", encoding="utf-8") as f:
        f.write(manifest)

    # -----------------------------
    # ZIP QTI PACKAGE
    # -----------------------------
    zip_path = os.path.join(job_dir, "qti_package.zip")
    zipf     = zipfile.ZipFile(zip_path, "w")

    for root, dirs, files in os.walk(job_dir):
        for file in files:
            if file.endswith((".zip", ".docx")):
                continue
            path    = os.path.join(root, file)
            arcname = os.path.relpath(path, job_dir)
            zipf.write(path, arcname)

    zipf.close()

    print(f"QTI 2.1 package created: {zip_path}")
    return zip_path
