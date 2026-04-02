"""
DOCX Question Parser
Parses Word documents following the formatting conventions:
  - Questions start with a number: "1 What is..."
  - MCQ options labeled A B C D
  - Answer key section begins with "ANSWERS"
  - Structured questions extracted as whole units
"""

import re
from docx import Document


def get_paragraph_text(para):
    """Extract full text from a paragraph, including all runs."""
    return para.text.strip()


def parse_docx(docx_path):
    """
    Parse a .docx file and return a list of question dicts.

    Each question dict has:
        number  : int
        text    : str  (full question body, may be multi-line)
        options : dict  e.g. {'A': 'Paris', 'B': 'London', ...}
        type    : 'mcq' | 'essay'
        answer  : str | None
    """
    doc = Document(docx_path)

    # Collect all non-empty paragraph texts (tables too, flattened)
    paragraphs = []
    for para in doc.paragraphs:
        text = get_paragraph_text(para)
        if text:
            paragraphs.append(text)

    # Also pull text from tables (some docs put options in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = get_paragraph_text(para)
                    if text:
                        paragraphs.append(text)

    questions = []
    answers = {}
    in_answers = False
    current_question = None

    # Patterns
    QUESTION_RE = re.compile(r'^(\d+)\s{1,4}(.+)')
    OPTION_RE   = re.compile(r'^([A-D])\s{1,4}(.+)', re.IGNORECASE)
    ANSWERS_RE  = re.compile(r'^ANSWERS\s*$', re.IGNORECASE)

    for para in paragraphs:
        # ── Answer section marker ──────────────────────────────────────────
        if ANSWERS_RE.match(para):
            in_answers = True
            if current_question:
                questions.append(current_question)
                current_question = None
            continue

        # ── Inside answer section ─────────────────────────────────────────
        if in_answers:
            # Expect lines like: "1 A"  or  "1 B,C"  or  "1 Paris"
            ans_match = re.match(r'^(\d+)\s+(.+)', para)
            if ans_match:
                q_num  = int(ans_match.group(1))
                answer = ans_match.group(2).strip()
                answers[q_num] = answer
            else:
                # Continuation answer (subsequent parts, no leading number)
                # Attach to the last seen question number
                if answers:
                    last_q = max(answers.keys())
                    answers[last_q] += '\n' + para
            continue

        # ── Question start ────────────────────────────────────────────────
        q_match = QUESTION_RE.match(para)
        if q_match:
            if current_question:
                questions.append(current_question)
            q_num  = int(q_match.group(1))
            q_text = q_match.group(2).strip()
            current_question = {
                'number':  q_num,
                'text':    q_text,
                'options': {},
                'type':    'essay',
                'answer':  None,
            }
            continue

        # ── MCQ option ────────────────────────────────────────────────────
        opt_match = OPTION_RE.match(para)
        if opt_match and current_question:
            opt_key  = opt_match.group(1).upper()
            opt_text = opt_match.group(2).strip()
            current_question['options'][opt_key] = opt_text
            current_question['type'] = 'mcq'
            continue

        # ── Continuation of current question ─────────────────────────────
        if current_question:
            current_question['text'] += '\n' + para

    # Flush last question
    if current_question:
        questions.append(current_question)

    # Attach answers
    for q in questions:
        q_num = q['number']
        if q_num in answers:
            q['answer'] = answers[q_num]

    return questions
